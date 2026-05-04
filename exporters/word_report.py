"""
Generate a full Word report (RTL Hebrew) for the HR Pulse dashboard.
"""
import io
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── RTL helpers ───────────────────────────────────────────────────────────────

def set_rtl_paragraph(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def add_rtl_paragraph(doc, text, bold=False, font_size=11, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    para = doc.add_paragraph()
    para.alignment = align
    set_rtl_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    return para


def add_rtl_heading(doc, text, level=1):
    para = doc.add_heading(level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(para)
    run = para.runs[0] if para.runs else para.add_run(text)
    if not para.runs:
        pass
    else:
        run.text = text
    run.font.name = 'Arial'
    set_rtl_run(run)
    return para


def set_cell_rtl(cell, text, bold=False, font_size=10):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    set_rtl_run(run)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ── Main report generator ────────────────────────────────────────────────────

def generate_word_report(stats: dict, dept_df, manager_df, employees_df) -> bytes:
    doc = Document()

    # Page setup: RTL
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set document-level RTL
    doc_settings = doc.settings.element
    bidi_elem = OxmlElement('w:bidi')
    doc_settings.append(bidi_elem)

    today = date.today().strftime("%d/%m/%Y")

    # ── Cover page ────────────────────────────────────────────────────────────
    add_rtl_paragraph(doc, "HR PULSE", bold=True, font_size=28,
                      color=(26, 115, 232), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_paragraph(doc, "דוח ניתוח עובדים", bold=True, font_size=22,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_paragraph(doc, f"תאריך: {today}", font_size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rtl_paragraph(doc, "סודי — לעיני הנהלה בלבד", font_size=11,
                      color=(234, 67, 53), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ── Executive summary ─────────────────────────────────────────────────────
    add_rtl_heading(doc, "תקציר מנהלים", level=1)

    high_risk_pct = stats.get('high_risk_pct', 0)
    avg_overall = stats.get('avg_overall', 0)
    total = stats.get('total_employees', 0)

    bullets = [
        f"סה\"כ {total} עובדים השתתפו בסקר שביעות רצון.",
        f"ציון מחוברות כולל: {avg_overall}/5 — {'מצב תקין' if avg_overall >= 3.5 else 'דורש תשומת לב' if avg_overall >= 3.0 else 'מצב דורש טיפול מיידי'}.",
        f"{stats.get('high_risk_count', 0)} עובדים ({high_risk_pct}%) מוגדרים בסיכון עזיבה גבוה.",
        f"eNPS ממוצע: {stats.get('avg_enps', 0)}.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(p)
        run = p.runs[0] if p.runs else p.add_run(b)
        run.text = b
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        set_rtl_run(run)

    doc.add_paragraph()

    # ── Key metrics table ─────────────────────────────────────────────────────
    add_rtl_heading(doc, "ממצאים עיקריים — מדדי מפתח", level=1)

    kpi_data = [
        ("מדד", "ערך", "הערכה"),
        ("סה\"כ עובדים", str(total), ""),
        ("ציון מחוברות כולל", f"{avg_overall}/5", "✅" if avg_overall >= 3.5 else "⚠️" if avg_overall >= 3.0 else "🔴"),
        ("eNPS ממוצע", str(stats.get('avg_enps', 0)), ""),
        ("% בסיכון עזיבה גבוה", f"{high_risk_pct}%", "✅" if high_risk_pct < 15 else "⚠️" if high_risk_pct < 25 else "🔴"),
        ("שביעות רצון שכר", f"{stats.get('avg_salary', 0)}/5", ""),
        ("שביעות רצון פיתוח", f"{stats.get('avg_development', 0)}/5", ""),
        ("שביעות רצון ממנהל", f"{stats.get('avg_manager', 0)}/5", ""),
        ("שביעות רצון מאיזון", f"{stats.get('avg_balance', 0)}/5", ""),
    ]

    table = doc.add_table(rows=len(kpi_data), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    for i, row_data in enumerate(kpi_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            set_cell_rtl(row.cells[j], cell_text, bold=(i == 0), font_size=10)
            if i == 0:
                set_cell_bg(row.cells[j], "1a73e8")
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # ── Department analysis ───────────────────────────────────────────────────
    add_rtl_heading(doc, "ניתוח מחלקות", level=1)
    add_rtl_paragraph(doc, "הטבלה הבאה מציגה את ציוני שביעות הרצון לפי מחלקה:", font_size=11)
    add_rtl_paragraph(doc, "גרפים זמינים בדשבורד האינטראקטיבי", font_size=10, color=(100, 100, 100))

    dept_cols = ["department", "employee_count", "avg_overall", "avg_salary",
                 "avg_development", "avg_manager", "avg_balance", "avg_enps", "high_risk_pct"]
    dept_headers = ["מחלקה", "עובדים", "ציון כולל", "שכר", "פיתוח", "מנהל", "איזון", "eNPS", "% סיכון גבוה"]

    if not dept_df.empty:
        avail_cols = [c for c in dept_cols if c in dept_df.columns]
        avail_headers = [dept_headers[dept_cols.index(c)] for c in avail_cols]
        table2 = doc.add_table(rows=len(dept_df) + 1, cols=len(avail_cols))
        table2.style = 'Table Grid'
        table2.alignment = WD_TABLE_ALIGNMENT.RIGHT

        # Header row
        for j, h in enumerate(avail_headers):
            set_cell_rtl(table2.rows[0].cells[j], h, bold=True, font_size=9)
            set_cell_bg(table2.rows[0].cells[j], "e8f0fe")

        for i, (_, row_data) in enumerate(dept_df[avail_cols].iterrows()):
            for j, val in enumerate(row_data):
                set_cell_rtl(table2.rows[i + 1].cells[j], str(val), font_size=9)

    doc.add_paragraph()

    # ── Manager analysis ──────────────────────────────────────────────────────
    add_rtl_heading(doc, "ניתוח מנהלים", level=1)

    if not manager_df.empty:
        mgr_sorted = manager_df.sort_values("engagement_score", ascending=False)

        add_rtl_paragraph(doc, "🏆 Top 5 מנהלים:", bold=True, font_size=12)
        top5 = mgr_sorted.head(5)
        for _, row_data in top5.iterrows():
            add_rtl_paragraph(doc,
                f"  {row_data['manager_name']} | ציון מעורבות: {row_data.get('engagement_score', 'N/A')} | עובדים: {row_data.get('employee_count', 'N/A')}",
                font_size=10)

        doc.add_paragraph()
        add_rtl_paragraph(doc, "⚠️ Bottom 5 מנהלים (טעונים שיפור):", bold=True, font_size=12)
        bot5 = mgr_sorted.tail(5)
        for _, row_data in bot5.iterrows():
            add_rtl_paragraph(doc,
                f"  {row_data['manager_name']} | ציון מעורבות: {row_data.get('engagement_score', 'N/A')} | % סיכון גבוה: {row_data.get('high_risk_pct', 'N/A')}%",
                font_size=10)

    doc.add_paragraph()

    # ── Flight risk analysis ──────────────────────────────────────────────────
    add_rtl_heading(doc, "ניתוח סיכון עזיבה", level=1)
    if not employees_df.empty:
        high_risk = employees_df[employees_df['risk_level'] == 'גבוה']
        mid_risk = employees_df[employees_df['risk_level'] == 'בינוני']
        low_risk = employees_df[employees_df['risk_level'] == 'נמוך']
        add_rtl_paragraph(doc, f"סיכון גבוה: {len(high_risk)} עובדים ({round(100*len(high_risk)/len(employees_df),1)}%)", font_size=11)
        add_rtl_paragraph(doc, f"סיכון בינוני: {len(mid_risk)} עובדים ({round(100*len(mid_risk)/len(employees_df),1)}%)", font_size=11)
        add_rtl_paragraph(doc, f"סיכון נמוך: {len(low_risk)} עובדים ({round(100*len(low_risk)/len(employees_df),1)}%)", font_size=11)

        if not dept_df.empty and 'high_risk_pct' in dept_df.columns:
            doc.add_paragraph()
            add_rtl_paragraph(doc, "מחלקות עם שיעור סיכון גבוה:", bold=True, font_size=11)
            worst_depts = dept_df.nlargest(3, 'high_risk_pct')
            for _, dr in worst_depts.iterrows():
                add_rtl_paragraph(doc, f"  {dr['department']}: {dr['high_risk_pct']}%", font_size=10)

    doc.add_paragraph()

    # ── Recommendations ───────────────────────────────────────────────────────
    add_rtl_heading(doc, "המלצות אופרטיביות", level=1)

    recs = _build_recommendations(stats, dept_df)
    for i, rec in enumerate(recs, 1):
        add_rtl_paragraph(doc, f"{i}. {rec['title']}", bold=True, font_size=11)
        add_rtl_paragraph(doc, f"   {rec['detail']}", font_size=10)
        add_rtl_paragraph(doc, f"   עדיפות: {rec['priority']} | השפעה משוערת: {rec['roi']}", font_size=10, color=(100, 100, 100))
        doc.add_paragraph()

    # ── 90-day action plan ────────────────────────────────────────────────────
    add_rtl_heading(doc, "תוכנית פעולה ל-90 יום", level=1)
    plan = [
        ("ימים 1-30", "זימון שיחות אישיות עם עובדים בסיכון גבוה. שיחת Coaching עם מנהלים בBottom 5."),
        ("ימים 31-60", "סקר שכר חיצוני. הקמת ועדת קידום. סדנאות מנהיגות ל-Bottom 5."),
        ("ימים 61-90", "הצגת מסלולי קידום ברורים. בדיקת יישום המלצות. סקר pulse מקוצר."),
    ]
    for period, action in plan:
        add_rtl_paragraph(doc, f"{period}:", bold=True, font_size=11)
        add_rtl_paragraph(doc, f"  {action}", font_size=10)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_recommendations(stats: dict, dept_df) -> list:
    recs = []
    if stats.get('avg_salary', 5) < 3.2:
        recs.append({
            "title": "שכר ותגמול — סקר שוק",
            "detail": "ביצוע סקר שכר חיצוני ופנימי. התאמת מסגרות שכר לשוק בתוך 60 יום.",
            "priority": "גבוה", "roi": "הפחתת עזיבות ב-15-20%"
        })
    if stats.get('high_risk_pct', 0) > 20:
        recs.append({
            "title": "שימור עובדים — תוכנית Retention",
            "detail": "הקמת תוכנית שימור ממוקדת לעובדים בסיכון גבוה: שיחות אישיות, תוכניות פיתוח, ומענקי שימור.",
            "priority": "גבוה", "roi": "חיסכון ממוצע 50-100K ₪ לכל עובד נשמר"
        })
    if stats.get('avg_development', 5) < 3.0:
        recs.append({
            "title": "פיתוח מקצועי — מסלולי קידום",
            "detail": "הגדרת מסלולי קידום ברורים בכל מחלקה. תוכנית הכשרות רלוונטיות. מנטורינג פנים-ארגוני.",
            "priority": "בינוני", "roi": "שיפור מחוברות ב-10-15%"
        })
    if stats.get('avg_manager', 5) < 3.2:
        recs.append({
            "title": "מנהיגות — פיתוח מנהלים",
            "detail": "הדרכות מנהיגות לכלל המנהלים. Coaching ממוקד לBottom 5. KPIs חדשים למנהלים הכוללים מחוברות צוות.",
            "priority": "גבוה", "roi": "שיפור ציון מנהל ב-0.3-0.5 נקודות"
        })
    if not dept_df.empty and 'high_risk_pct' in dept_df.columns:
        worst = dept_df.nlargest(1, 'high_risk_pct').iloc[0]
        if worst['high_risk_pct'] > stats.get('high_risk_pct', 0) * 1.5:
            recs.append({
                "title": f"מחלקת {worst['department']} — התייחסות מיידית",
                "detail": f"שיעור סיכון עזיבה של {worst['high_risk_pct']}% — פי 1.5 מהממוצע. נדרשת תוכנית מיוחדת.",
                "priority": "גבוה", "roi": "מניעת גל עזיבות מרוכז"
            })
    if not recs:
        recs.append({
            "title": "שמירה על מגמה חיובית",
            "detail": "המדדים תקינים. המשך מעקב רבעוני. חיזוק נקודות חוזק ארגוניות.",
            "priority": "נמוך", "roi": "שמירה על רמת מחוברות קיימת"
        })
    return recs
