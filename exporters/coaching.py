"""
Generate coaching letters for Bottom 5 managers as a single Word document (RTL Hebrew).
Each manager gets a page break.
"""
import io
import random
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_rtl_paragraph(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)


def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def add_rtl_para(doc, text, bold=False, font_size=11, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    para = doc.add_paragraph()
    para.alignment = align
    set_rtl_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_rtl_run(run)
    return para


def _get_recommendations(scores: dict) -> list:
    recs = []
    sorted_scores = sorted(scores.items(), key=lambda x: x[1])
    lowest_dim, lowest_val = sorted_scores[0]

    dim_to_rec = {
        "שביעות רצון ממנהל": "לקיים פגישות אחד-על-אחד שבועיות עם כל חבר צוות. לספק פידבק בונה ומעודד באופן שוטף.",
        "שביעות רצון מפיתוח": "לשוחח עם כל עובד על מסלול הקריירה שלו ולהגדיר יעדי פיתוח אישיים לחצי שנה הקרובה.",
        "שביעות רצון מאיזון עבודה-חיים": "לבחון עומס עבודה בצוות ולהקצות משאבים נוספים אם נדרש. לעודד גבולות בריאים בין עבודה לחיים.",
        "שביעות רצון שכר": "לבצע בדיקת שוק ולהעלות נושא השכר עם משאבי אנוש עבור עובדים שמשכורתם נמוכה מהשוק.",
    }

    for dim, val in sorted_scores[:2]:
        rec = dim_to_rec.get(dim, "לקיים שיחות פתוחות עם הצוות ולהקשיב לצרכים שלהם.")
        recs.append(f"שיפור {dim} ({val:.1f}/5): {rec}")

    return recs


def generate_coaching_letters(manager_df, employees_df) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # Document-level RTL
    doc_settings = doc.settings.element
    bidi_elem = OxmlElement('w:bidi')
    doc_settings.append(bidi_elem)

    today = date.today().strftime("%d/%m/%Y")

    # Get bottom 5 managers by engagement score
    if manager_df.empty:
        add_rtl_para(doc, "אין נתוני מנהלים זמינים.", font_size=12)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    bottom5 = manager_df.sort_values("engagement_score", ascending=True).head(5)

    for idx, (_, mgr_row) in enumerate(bottom5.iterrows()):
        if idx > 0:
            doc.add_page_break()

        manager_name = mgr_row['manager_name']
        dept = mgr_row.get('department', '')
        emp_count = int(mgr_row.get('employee_count', 0))

        # Get employees under this manager
        mgr_employees = employees_df[employees_df['manager_name'] == manager_name] if not employees_df.empty else None

        # Scores
        mgr_sat = float(mgr_row.get('avg_manager', 0))
        dev_sat = float(mgr_row.get('avg_development', 0))
        bal_sat = float(mgr_row.get('avg_balance', 0))
        sal_sat = float(mgr_row.get('avg_salary', 0))
        avg_enps = float(mgr_row.get('avg_enps', 0))
        high_risk_pct = float(mgr_row.get('high_risk_pct', 0))

        # ── Letter header ──────────────────────────────────────────────────
        add_rtl_para(doc, f"תאריך: {today}", font_size=10, color=(120, 120, 120))
        add_rtl_para(doc, f"לכבוד {manager_name},", bold=True, font_size=13)
        doc.add_paragraph()

        add_rtl_para(doc,
            f"בעקבות סקר שביעות רצון העובדים, ברצוננו לשתף אותך בממצאים מצוות {dept} ({emp_count} עובדים).",
            font_size=11)
        doc.add_paragraph()

        # ── Team scores ────────────────────────────────────────────────────
        add_rtl_para(doc, "ציוני הצוות שלך:", bold=True, font_size=11)

        score_lines = [
            f"• שביעות רצון ממנהל: {mgr_sat:.1f}/5",
            f"• שביעות רצון מפיתוח: {dev_sat:.1f}/5",
            f"• שביעות רצון מאיזון עבודה-חיים: {bal_sat:.1f}/5",
            f"• שביעות רצון שכר: {sal_sat:.1f}/5",
            f"• eNPS ממוצע: {avg_enps:.0f}",
            f"• % בסיכון עזיבה גבוה: {high_risk_pct:.0f}%",
        ]
        for line in score_lines:
            add_rtl_para(doc, line, font_size=11)

        doc.add_paragraph()

        # ── Quotes from team ───────────────────────────────────────────────
        add_rtl_para(doc, "נושאים עיקריים שעלו ממשוב הצוות:", bold=True, font_size=11)

        quotes = []
        if mgr_employees is not None and not mgr_employees.empty and 'open_feedback' in mgr_employees.columns:
            valid_feedback = mgr_employees['open_feedback'].dropna().tolist()
            n = min(4, len(valid_feedback))
            if n > 0:
                quotes = random.sample(valid_feedback, n)

        if quotes:
            for q in quotes:
                add_rtl_para(doc, f'  "...{q}..."', font_size=10, color=(80, 80, 80))
        else:
            add_rtl_para(doc, "  לא נמצא משוב פתוח לצוות זה.", font_size=10, color=(120, 120, 120))

        doc.add_paragraph()

        # ── Strengths ──────────────────────────────────────────────────────
        scores_dict = {
            "שביעות רצון ממנהל": mgr_sat,
            "שביעות רצון מפיתוח": dev_sat,
            "שביעות רצון מאיזון עבודה-חיים": bal_sat,
            "שביעות רצון שכר": sal_sat,
        }
        strengths = [(k, v) for k, v in scores_dict.items() if v >= 3.5]
        if strengths:
            add_rtl_para(doc, "נקודות חוזק:", bold=True, font_size=11)
            for k, v in strengths:
                add_rtl_para(doc, f"  ✅ {k}: {v:.1f}/5 — ממשיך לעשות עבודה טובה!", font_size=10)
            doc.add_paragraph()

        # ── Areas for improvement ──────────────────────────────────────────
        weak = [(k, v) for k, v in scores_dict.items() if v < 3.5]
        weak.sort(key=lambda x: x[1])
        if weak:
            add_rtl_para(doc, "תחומים לשיפור עיקריים:", bold=True, font_size=11)
            for k, v in weak[:3]:
                add_rtl_para(doc, f"  ⚠️ {k}: {v:.1f}/5", font_size=10)
            doc.add_paragraph()

        # ── Recommendations ────────────────────────────────────────────────
        add_rtl_para(doc, "המלצות לפעולה:", bold=True, font_size=11)

        recs = _get_recommendations(scores_dict)
        for i, rec in enumerate(recs, 1):
            add_rtl_para(doc, f"  {i}. {rec}", font_size=11)

        if high_risk_pct > 25:
            add_rtl_para(doc,
                f"  {len(recs)+1}. שיעור סיכון עזיבה של {high_risk_pct:.0f}% גבוה מאוד — לקיים שיחות שימור אישיות עם עובדים בסיכון מיידי.",
                font_size=11)
        add_rtl_para(doc,
            "  • להשתתף בסדנת מנהיגות המוצעת על-ידי צוות פיתוח הארגון.",
            font_size=11)

        doc.add_paragraph()

        # ── Closing ────────────────────────────────────────────────────────
        add_rtl_para(doc,
            "אנו מאמינים ביכולתך לחולל שינוי חיובי בצוות. אנו כאן לתמוך בך בתהליך.",
            font_size=11)
        doc.add_paragraph()
        add_rtl_para(doc, "בכבוד רב,", font_size=11)
        add_rtl_para(doc, "צוות משאבי אנוש", bold=True, font_size=11)
        add_rtl_para(doc, "HR Pulse", font_size=10, color=(26, 115, 232))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
